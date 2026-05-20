from pathlib import Path


def render_docx_report(context: dict, output_path: str | Path = "reports/outputs/relatorio_flyrock.docx") -> Path | None:
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(context["config"]["project"]["name"], 0)
        doc.add_paragraph("Ferramenta de analise; nao substitui responsavel tecnico habilitado.")
        for title in ["Objetivo", "Metodologia", "Calibracao do modelo", "Raios de seguranca", "Limitacoes"]:
            doc.add_heading(title, level=1)
            doc.add_paragraph("Conteudo tecnico conforme relatorio HTML gerado pelo sistema.")
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
        return output
    except Exception:
        return None
